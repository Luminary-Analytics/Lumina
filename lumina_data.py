#!/usr/bin/env python3
"""
╔═══════════════════════════════════════════════════════════════════════════════╗
║                         LUMINA DATA SYSTEM                                     ║
║                                                                               ║
║  Database operations and document handling for Lumina.                        ║
║  Enables creation of databases, CRUD operations, and document read/write.    ║
║                                                                               ║
║  Capabilities:                                                                 ║
║  - SQLite database creation and management                                    ║
║  - PDF reading and creation                                                   ║
║  - Word document reading and creation                                         ║
║  - Spreadsheet operations                                                     ║
║  - E-book reading (EPUB)                                                      ║
║                                                                               ║
║  Created: 2025-12-07                                                          ║
╚═══════════════════════════════════════════════════════════════════════════════╝
"""

import os
import sys
import json
import sqlite3
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, field

# ═══════════════════════════════════════════════════════════════════════════════
# DATABASE MANAGER
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class TableSchema:
    """Schema for a database table."""
    name: str
    columns: List[Dict[str, str]]  # [{"name": "id", "type": "INTEGER", "constraints": "PRIMARY KEY"}]
    
    def to_sql(self) -> str:
        """Generate CREATE TABLE SQL."""
        col_defs = []
        for col in self.columns:
            col_def = f"{col['name']} {col['type']}"
            if col.get('constraints'):
                col_def += f" {col['constraints']}"
            col_defs.append(col_def)
        
        return f"CREATE TABLE IF NOT EXISTS {self.name} ({', '.join(col_defs)})"


class DatabaseManager:
    """Manages SQLite databases for Lumina."""
    
    def __init__(self, workspace_path: Path):
        self.workspace_path = workspace_path
        self.databases_path = workspace_path / "databases"
        self.databases_path.mkdir(parents=True, exist_ok=True)
        
        self.active_connections: Dict[str, sqlite3.Connection] = {}
    
    def create_database(self, name: str, schemas: List[TableSchema] = None) -> bool:
        """Create a new database with optional initial schemas."""
        db_path = self.databases_path / f"{name}.db"
        
        try:
            conn = sqlite3.connect(str(db_path))
            conn.execute("PRAGMA journal_mode=WAL")
            
            if schemas:
                for schema in schemas:
                    conn.execute(schema.to_sql())
                conn.commit()
            
            conn.close()
            return True
        except Exception as e:
            print(f"Error creating database: {e}")
            return False
    
    def get_connection(self, name: str) -> Optional[sqlite3.Connection]:
        """Get a connection to a database."""
        if name in self.active_connections:
            return self.active_connections[name]
        
        db_path = self.databases_path / f"{name}.db"
        if not db_path.exists():
            return None
        
        try:
            conn = sqlite3.connect(str(db_path))
            conn.row_factory = sqlite3.Row
            self.active_connections[name] = conn
            return conn
        except:
            return None
    
    def close_connection(self, name: str) -> None:
        """Close a database connection."""
        if name in self.active_connections:
            self.active_connections[name].close()
            del self.active_connections[name]
    
    def list_databases(self) -> List[str]:
        """List all available databases."""
        return [f.stem for f in self.databases_path.glob("*.db")]
    
    def get_tables(self, db_name: str) -> List[str]:
        """Get list of tables in a database."""
        conn = self.get_connection(db_name)
        if not conn:
            return []
        
        cursor = conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' ORDER BY name"
        )
        return [row[0] for row in cursor.fetchall()]
    
    def get_table_info(self, db_name: str, table_name: str) -> List[Dict]:
        """Get column information for a table."""
        conn = self.get_connection(db_name)
        if not conn:
            return []
        
        cursor = conn.execute(f"PRAGMA table_info({table_name})")
        return [
            {"cid": row[0], "name": row[1], "type": row[2], "notnull": row[3], "pk": row[5]}
            for row in cursor.fetchall()
        ]
    
    def execute(self, db_name: str, sql: str, params: tuple = ()) -> Optional[List[Dict]]:
        """Execute SQL and return results."""
        conn = self.get_connection(db_name)
        if not conn:
            return None
        
        try:
            cursor = conn.execute(sql, params)
            
            # Check if it's a SELECT
            if sql.strip().upper().startswith("SELECT"):
                columns = [desc[0] for desc in cursor.description]
                return [dict(zip(columns, row)) for row in cursor.fetchall()]
            else:
                conn.commit()
                return [{"rows_affected": cursor.rowcount}]
        except Exception as e:
            return [{"error": str(e)}]
    
    def insert(self, db_name: str, table: str, data: Dict) -> bool:
        """Insert a row into a table."""
        conn = self.get_connection(db_name)
        if not conn:
            return False
        
        columns = ', '.join(data.keys())
        placeholders = ', '.join(['?' for _ in data])
        sql = f"INSERT INTO {table} ({columns}) VALUES ({placeholders})"
        
        try:
            conn.execute(sql, tuple(data.values()))
            conn.commit()
            return True
        except:
            return False
    
    def insert_many(self, db_name: str, table: str, data_list: List[Dict]) -> int:
        """Insert multiple rows."""
        if not data_list:
            return 0
        
        conn = self.get_connection(db_name)
        if not conn:
            return 0
        
        columns = ', '.join(data_list[0].keys())
        placeholders = ', '.join(['?' for _ in data_list[0]])
        sql = f"INSERT INTO {table} ({columns}) VALUES ({placeholders})"
        
        try:
            cursor = conn.executemany(sql, [tuple(d.values()) for d in data_list])
            conn.commit()
            return cursor.rowcount
        except:
            return 0
    
    def update(self, db_name: str, table: str, data: Dict, where: str, params: tuple = ()) -> bool:
        """Update rows in a table."""
        conn = self.get_connection(db_name)
        if not conn:
            return False
        
        set_clause = ', '.join([f"{k} = ?" for k in data.keys()])
        sql = f"UPDATE {table} SET {set_clause} WHERE {where}"
        
        try:
            conn.execute(sql, tuple(data.values()) + params)
            conn.commit()
            return True
        except:
            return False
    
    def delete(self, db_name: str, table: str, where: str, params: tuple = ()) -> int:
        """Delete rows from a table."""
        conn = self.get_connection(db_name)
        if not conn:
            return 0
        
        sql = f"DELETE FROM {table} WHERE {where}"
        
        try:
            cursor = conn.execute(sql, params)
            conn.commit()
            return cursor.rowcount
        except:
            return 0
    
    def query(self, db_name: str, sql: str, params: tuple = ()) -> List[Dict]:
        """Execute a SELECT query."""
        result = self.execute(db_name, sql, params)
        return result if result and "error" not in result[0] else []


# ═══════════════════════════════════════════════════════════════════════════════
# KNOWLEDGE BASE - Specialized database for facts and knowledge
# ═══════════════════════════════════════════════════════════════════════════════

class KnowledgeBase:
    """A specialized database for Lumina's knowledge and facts."""
    
    def __init__(self, db_manager: DatabaseManager):
        self.db = db_manager
        self.db_name = "knowledge"
        self._ensure_schema()
    
    def _ensure_schema(self):
        """Ensure the knowledge database exists with proper schema."""
        schemas = [
            TableSchema("facts", [
                {"name": "id", "type": "INTEGER", "constraints": "PRIMARY KEY AUTOINCREMENT"},
                {"name": "content", "type": "TEXT", "constraints": "NOT NULL"},
                {"name": "category", "type": "TEXT"},
                {"name": "source", "type": "TEXT"},
                {"name": "confidence", "type": "REAL", "constraints": "DEFAULT 0.8"},
                {"name": "created_at", "type": "TEXT"},
                {"name": "last_accessed", "type": "TEXT"},
                {"name": "access_count", "type": "INTEGER", "constraints": "DEFAULT 0"}
            ]),
            TableSchema("topics", [
                {"name": "id", "type": "INTEGER", "constraints": "PRIMARY KEY AUTOINCREMENT"},
                {"name": "name", "type": "TEXT", "constraints": "UNIQUE NOT NULL"},
                {"name": "description", "type": "TEXT"},
                {"name": "parent_id", "type": "INTEGER"}
            ]),
            TableSchema("fact_topics", [
                {"name": "fact_id", "type": "INTEGER"},
                {"name": "topic_id", "type": "INTEGER"},
                {"name": "PRIMARY KEY", "type": "(fact_id, topic_id)", "constraints": ""}
            ]),
            TableSchema("insights", [
                {"name": "id", "type": "INTEGER", "constraints": "PRIMARY KEY AUTOINCREMENT"},
                {"name": "content", "type": "TEXT", "constraints": "NOT NULL"},
                {"name": "derived_from", "type": "TEXT"},  # JSON list of fact IDs
                {"name": "created_at", "type": "TEXT"}
            ])
        ]
        
        self.db.create_database(self.db_name, schemas)
    
    def add_fact(self, content: str, category: str = None, source: str = None, 
                 confidence: float = 0.8, topics: List[str] = None) -> int:
        """Add a fact to the knowledge base."""
        data = {
            "content": content,
            "category": category,
            "source": source,
            "confidence": confidence,
            "created_at": datetime.now().isoformat(),
            "last_accessed": datetime.now().isoformat()
        }
        
        if self.db.insert(self.db_name, "facts", data):
            # Get the inserted ID
            result = self.db.query(self.db_name, "SELECT last_insert_rowid() as id")
            if result:
                fact_id = result[0]["id"]
                
                # Add topics
                if topics:
                    for topic in topics:
                        self._add_topic_to_fact(fact_id, topic)
                
                return fact_id
        return -1
    
    def _add_topic_to_fact(self, fact_id: int, topic_name: str) -> None:
        """Link a fact to a topic, creating the topic if needed."""
        # Ensure topic exists
        result = self.db.query(self.db_name, 
            "SELECT id FROM topics WHERE name = ?", (topic_name,))
        
        if result:
            topic_id = result[0]["id"]
        else:
            self.db.insert(self.db_name, "topics", {"name": topic_name})
            result = self.db.query(self.db_name, "SELECT last_insert_rowid() as id")
            topic_id = result[0]["id"]
        
        # Link fact to topic
        self.db.insert(self.db_name, "fact_topics", 
            {"fact_id": fact_id, "topic_id": topic_id})
    
    def get_facts(self, category: str = None, topic: str = None, 
                  limit: int = 50) -> List[Dict]:
        """Get facts, optionally filtered."""
        if topic:
            sql = """
                SELECT f.* FROM facts f
                JOIN fact_topics ft ON f.id = ft.fact_id
                JOIN topics t ON ft.topic_id = t.id
                WHERE t.name = ?
                ORDER BY f.created_at DESC
                LIMIT ?
            """
            return self.db.query(self.db_name, sql, (topic, limit))
        elif category:
            sql = "SELECT * FROM facts WHERE category = ? ORDER BY created_at DESC LIMIT ?"
            return self.db.query(self.db_name, sql, (category, limit))
        else:
            sql = "SELECT * FROM facts ORDER BY created_at DESC LIMIT ?"
            return self.db.query(self.db_name, sql, (limit,))
    
    def search_facts(self, query: str, limit: int = 20) -> List[Dict]:
        """Search facts by content."""
        sql = "SELECT * FROM facts WHERE content LIKE ? ORDER BY confidence DESC LIMIT ?"
        return self.db.query(self.db_name, sql, (f"%{query}%", limit))
    
    def get_topics(self) -> List[Dict]:
        """Get all topics with fact counts."""
        sql = """
            SELECT t.*, COUNT(ft.fact_id) as fact_count
            FROM topics t
            LEFT JOIN fact_topics ft ON t.id = ft.topic_id
            GROUP BY t.id
            ORDER BY fact_count DESC
        """
        return self.db.query(self.db_name, sql)
    
    def add_insight(self, content: str, derived_from: List[int] = None) -> int:
        """Add an insight derived from facts."""
        data = {
            "content": content,
            "derived_from": json.dumps(derived_from or []),
            "created_at": datetime.now().isoformat()
        }
        
        if self.db.insert(self.db_name, "insights", data):
            result = self.db.query(self.db_name, "SELECT last_insert_rowid() as id")
            return result[0]["id"] if result else -1
        return -1
    
    def get_stats(self) -> Dict:
        """Get knowledge base statistics."""
        facts = self.db.query(self.db_name, "SELECT COUNT(*) as count FROM facts")
        topics = self.db.query(self.db_name, "SELECT COUNT(*) as count FROM topics")
        insights = self.db.query(self.db_name, "SELECT COUNT(*) as count FROM insights")
        
        return {
            "total_facts": facts[0]["count"] if facts else 0,
            "total_topics": topics[0]["count"] if topics else 0,
            "total_insights": insights[0]["count"] if insights else 0
        }


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT READER
# ═══════════════════════════════════════════════════════════════════════════════

class DocumentReader:
    """Reads various document formats."""
    
    def __init__(self, workspace_path: Path):
        self.workspace_path = workspace_path
        self.documents_path = workspace_path / "documents"
        self.documents_path.mkdir(parents=True, exist_ok=True)
    
    def read_pdf(self, path: Union[str, Path]) -> Optional[str]:
        """Read text from a PDF file."""
        try:
            import PyPDF2
            
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = []
                for page in reader.pages:
                    text.append(page.extract_text())
                return '\n'.join(text)
        except ImportError:
            try:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    text = []
                    for page in pdf.pages:
                        text.append(page.extract_text() or '')
                    return '\n'.join(text)
            except ImportError:
                return None
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return None
    
    def read_word(self, path: Union[str, Path]) -> Optional[str]:
        """Read text from a Word document."""
        try:
            from docx import Document
            
            doc = Document(path)
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            return '\n'.join(text)
        except ImportError:
            return None
        except Exception as e:
            print(f"Error reading Word doc: {e}")
            return None
    
    def read_epub(self, path: Union[str, Path]) -> Optional[str]:
        """Read text from an EPUB e-book."""
        try:
            from ebooklib import epub
            from bs4 import BeautifulSoup
            
            book = epub.read_epub(path)
            text = []
            
            for item in book.get_items():
                if item.get_type() == 9:  # ITEM_DOCUMENT
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text.append(soup.get_text())
            
            return '\n'.join(text)
        except ImportError:
            return None
        except Exception as e:
            print(f"Error reading EPUB: {e}")
            return None
    
    def read_text(self, path: Union[str, Path]) -> Optional[str]:
        """Read a plain text file."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading text file: {e}")
            return None
    
    def read_spreadsheet(self, path: Union[str, Path], sheet_name: str = None) -> Optional[List[Dict]]:
        """Read data from a spreadsheet."""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(path)
            ws = wb[sheet_name] if sheet_name else wb.active
            
            data = []
            headers = [cell.value for cell in ws[1]]
            
            for row in ws.iter_rows(min_row=2, values_only=True):
                data.append(dict(zip(headers, row)))
            
            return data
        except ImportError:
            return None
        except Exception as e:
            print(f"Error reading spreadsheet: {e}")
            return None
    
    def read_auto(self, path: Union[str, Path]) -> Optional[str]:
        """Automatically detect format and read."""
        path = Path(path)
        ext = path.suffix.lower()
        
        if ext == '.pdf':
            return self.read_pdf(path)
        elif ext in ['.docx', '.doc']:
            return self.read_word(path)
        elif ext == '.epub':
            return self.read_epub(path)
        elif ext in ['.txt', '.md', '.py', '.json', '.csv']:
            return self.read_text(path)
        elif ext in ['.xlsx', '.xls']:
            data = self.read_spreadsheet(path)
            return json.dumps(data, indent=2) if data else None
        else:
            return self.read_text(path)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT WRITER
# ═══════════════════════════════════════════════════════════════════════════════

class DocumentWriter:
    """Creates various document formats."""
    
    def __init__(self, workspace_path: Path):
        self.workspace_path = workspace_path
        self.documents_path = workspace_path / "documents"
        self.documents_path.mkdir(parents=True, exist_ok=True)
    
    def write_pdf(self, filename: str, content: str, title: str = None) -> Optional[Path]:
        """Create a PDF document."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch
            
            path = self.documents_path / filename
            c = canvas.Canvas(str(path), pagesize=letter)
            width, height = letter
            
            # Title
            if title:
                c.setFont("Helvetica-Bold", 18)
                c.drawString(1*inch, height - 1*inch, title)
                y = height - 1.5*inch
            else:
                y = height - 1*inch
            
            # Content
            c.setFont("Helvetica", 11)
            lines = content.split('\n')
            
            for line in lines:
                if y < 1*inch:
                    c.showPage()
                    c.setFont("Helvetica", 11)
                    y = height - 1*inch
                
                # Wrap long lines
                while len(line) > 80:
                    c.drawString(1*inch, y, line[:80])
                    line = line[80:]
                    y -= 14
                    if y < 1*inch:
                        c.showPage()
                        c.setFont("Helvetica", 11)
                        y = height - 1*inch
                
                c.drawString(1*inch, y, line)
                y -= 14
            
            c.save()
            return path
        except ImportError:
            try:
                from fpdf import FPDF
                
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=11)
                
                if title:
                    pdf.set_font("Arial", 'B', 16)
                    pdf.cell(200, 10, txt=title, ln=True, align='C')
                    pdf.set_font("Arial", size=11)
                
                for line in content.split('\n'):
                    pdf.multi_cell(0, 10, txt=line)
                
                path = self.documents_path / filename
                pdf.output(str(path))
                return path
            except ImportError:
                return None
        except Exception as e:
            print(f"Error creating PDF: {e}")
            return None
    
    def write_word(self, filename: str, content: str, title: str = None) -> Optional[Path]:
        """Create a Word document."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            
            doc = Document()
            
            if title:
                doc.add_heading(title, 0)
            
            for para in content.split('\n\n'):
                doc.add_paragraph(para)
            
            path = self.documents_path / filename
            doc.save(str(path))
            return path
        except ImportError:
            return None
        except Exception as e:
            print(f"Error creating Word doc: {e}")
            return None
    
    def write_spreadsheet(self, filename: str, data: List[Dict], 
                          sheet_name: str = "Sheet1") -> Optional[Path]:
        """Create a spreadsheet."""
        try:
            from openpyxl import Workbook
            
            wb = Workbook()
            ws = wb.active
            ws.title = sheet_name
            
            if data:
                # Headers
                headers = list(data[0].keys())
                for col, header in enumerate(headers, 1):
                    ws.cell(row=1, column=col, value=header)
                
                # Data
                for row_num, row_data in enumerate(data, 2):
                    for col, header in enumerate(headers, 1):
                        ws.cell(row=row_num, column=col, value=row_data.get(header))
            
            path = self.documents_path / filename
            wb.save(str(path))
            return path
        except ImportError:
            return None
        except Exception as e:
            print(f"Error creating spreadsheet: {e}")
            return None
    
    def write_markdown(self, filename: str, content: str) -> Optional[Path]:
        """Create a Markdown file."""
        try:
            path = self.documents_path / filename
            with open(path, 'w', encoding='utf-8') as f:
                f.write(content)
            return path
        except Exception as e:
            print(f"Error creating Markdown: {e}")
            return None
    
    def write_html(self, filename: str, content: str, title: str = None) -> Optional[Path]:
        """Create an HTML file."""
        try:
            html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title or 'Document'}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #333; }}
    </style>
</head>
<body>
    {f'<h1>{title}</h1>' if title else ''}
    {content}
</body>
</html>"""
            
            path = self.documents_path / filename
            with open(path, 'w', encoding='utf-8') as f:
                f.write(html)
            return path
        except Exception as e:
            print(f"Error creating HTML: {e}")
            return None


# ═══════════════════════════════════════════════════════════════════════════════
# LUMINA DATA INTERFACE
# ═══════════════════════════════════════════════════════════════════════════════

class LuminaData:
    """Lumina's unified interface to data and documents."""
    
    def __init__(self, workspace_path: Path):
        self.workspace_path = workspace_path
        self.db_manager = DatabaseManager(workspace_path)
        self.knowledge = KnowledgeBase(self.db_manager)
        self.reader = DocumentReader(workspace_path)
        self.writer = DocumentWriter(workspace_path)
    
    # Database shortcuts
    def create_db(self, name: str, schemas: List[TableSchema] = None) -> bool:
        return self.db_manager.create_database(name, schemas)
    
    def query(self, db_name: str, sql: str, params: tuple = ()) -> List[Dict]:
        return self.db_manager.query(db_name, sql, params)
    
    def insert(self, db_name: str, table: str, data: Dict) -> bool:
        return self.db_manager.insert(db_name, table, data)
    
    # Knowledge shortcuts
    def learn_fact(self, content: str, category: str = None, source: str = None) -> int:
        return self.knowledge.add_fact(content, category, source)
    
    def recall_facts(self, query: str = None, category: str = None) -> List[Dict]:
        if query:
            return self.knowledge.search_facts(query)
        return self.knowledge.get_facts(category=category)
    
    # Document shortcuts
    def read_document(self, path: str) -> Optional[str]:
        return self.reader.read_auto(path)
    
    def create_pdf(self, filename: str, content: str, title: str = None) -> Optional[Path]:
        return self.writer.write_pdf(filename, content, title)
    
    def create_word(self, filename: str, content: str, title: str = None) -> Optional[Path]:
        return self.writer.write_word(filename, content, title)
    
    def create_report(self, title: str, content: str, format: str = "pdf") -> Optional[Path]:
        """Create a report in the specified format."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format == "pdf":
            return self.create_pdf(f"report_{timestamp}.pdf", content, title)
        elif format == "word":
            return self.create_word(f"report_{timestamp}.docx", content, title)
        elif format == "md":
            full_content = f"# {title}\n\n{content}"
            return self.writer.write_markdown(f"report_{timestamp}.md", full_content)
        elif format == "html":
            return self.writer.write_html(f"report_{timestamp}.html", content, title)
        return None
    
    def get_stats(self) -> Dict:
        """Get data system statistics."""
        return {
            "databases": self.db_manager.list_databases(),
            "knowledge": self.knowledge.get_stats()
        }


# ═══════════════════════════════════════════════════════════════════════════════
# INITIALIZATION
# ═══════════════════════════════════════════════════════════════════════════════

def initialize_data_system(workspace_path: Path) -> LuminaData:
    """Initialize Lumina's data system."""
    return LuminaData(workspace_path)


if __name__ == "__main__":
    # Test the system
    workspace = Path("lumina_workspace")
    workspace.mkdir(exist_ok=True)
    
    data = initialize_data_system(workspace)
    
    print("Data System Initialized!")
    print(f"Stats: {data.get_stats()}")
    
    # Test adding a fact
    fact_id = data.learn_fact(
        "The universe is approximately 13.8 billion years old",
        category="science",
        source="cosmology"
    )
    print(f"Added fact with ID: {fact_id}")

